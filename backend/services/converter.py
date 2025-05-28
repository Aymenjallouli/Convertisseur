from pathlib import Path
import asyncio
from typing import Optional
import pandas as pd
from docx import Document
import PyPDF2
from PIL import Image
from pptx import Presentation
import openpyxl
from pdf2docx import Converter
import img2pdf
import os

class FileConverter:
    def __init__(self):
        self.supported_conversions = {
            'docx': ['pdf', 'txt'],
            'pdf': ['txt', 'docx'],
            'xlsx': ['csv', 'txt'],
            'csv': ['xlsx', 'txt'],
            'txt': ['pdf', 'docx'],
            'pptx': ['pdf', 'txt'],
            'jpg': ['png', 'pdf', 'jpeg'],
            'jpeg': ['png', 'pdf', 'jpg'],
            'png': ['jpg', 'pdf', 'jpeg'],
            'bmp': ['png', 'jpg', 'jpeg', 'pdf']
        }
    
    async def convert_file(self, source_path: Path, target_path: Path, target_format: str) -> bool:
        """Convertit un fichier vers le format cible"""
        try:
            source_format = source_path.suffix[1:].lower()
            
            if source_format not in self.supported_conversions:
                return False
            
            if target_format not in self.supported_conversions[source_format]:
                return False
            
            # Appeler la méthode de conversion appropriée
            conversion_method = f"_convert_{source_format}_to_{target_format}"
            if hasattr(self, conversion_method):
                method = getattr(self, conversion_method)
                return await method(source_path, target_path)
            else:
                return False
                
        except Exception as e:
            print(f"Erreur lors de la conversion: {e}")
            return False
    
    # Conversions DOCX
    async def _convert_docx_to_txt(self, source_path: Path, target_path: Path) -> bool:
        try:
            doc = Document(source_path)
            text_content = []
            
            for paragraph in doc.paragraphs:
                text_content.append(paragraph.text)
            
            with open(target_path, 'w', encoding='utf-8') as f:
                f.write('\n'.join(text_content))
            
            return True
        except Exception as e:
            print(f"Erreur conversion DOCX vers TXT: {e}")
            return False

    async def _convert_docx_to_pdf(self, source_path: Path, target_path: Path) -> bool:
        try:
            from docx2pdf import convert
            convert(str(source_path), str(target_path))
            return True
        except Exception as e:
            print(f"Erreur conversion DOCX vers PDF: {e}")
            return False
    
    # Conversions PDF
    async def _convert_pdf_to_txt(self, source_path: Path, target_path: Path) -> bool:
        try:
            with open(source_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text_content = []
                
                for page in pdf_reader.pages:
                    text_content.append(page.extract_text())
            
            with open(target_path, 'w', encoding='utf-8') as f:
                f.write('\n'.join(text_content))
            
            return True
        except Exception as e:
            print(f"Erreur conversion PDF vers TXT: {e}")
            return False
    
    async def _convert_pdf_to_docx(self, source_path: Path, target_path: Path) -> bool:
        try:
            # Utiliser pdf2docx pour la conversion
            cv = Converter(str(source_path))
            cv.convert(str(target_path))
            cv.close()
            return True
        except Exception as e:
            print(f"Erreur conversion PDF vers DOCX: {e}")
            return False
    
    # Conversions Excel
    async def _convert_xlsx_to_csv(self, source_path: Path, target_path: Path) -> bool:
        try:
            df = pd.read_excel(source_path)
            df.to_csv(target_path, index=False, encoding='utf-8')
            return True
        except Exception as e:
            print(f"Erreur conversion XLSX vers CSV: {e}")
            return False
    
    async def _convert_xlsx_to_txt(self, source_path: Path, target_path: Path) -> bool:
        try:
            df = pd.read_excel(source_path)
            df.to_csv(target_path, index=False, encoding='utf-8', sep='\t')
            return True
        except Exception as e:
            print(f"Erreur conversion XLSX vers TXT: {e}")
            return False
    
    # Conversions CSV
    async def _convert_csv_to_xlsx(self, source_path: Path, target_path: Path) -> bool:
        try:
            df = pd.read_csv(source_path)
            df.to_excel(target_path, index=False)
            return True
        except Exception as e:
            print(f"Erreur conversion CSV vers XLSX: {e}")
            return False
    
    async def _convert_csv_to_txt(self, source_path: Path, target_path: Path) -> bool:
        try:
            df = pd.read_csv(source_path)
            df.to_csv(target_path, index=False, encoding='utf-8', sep='\t')
            return True
        except Exception as e:
            print(f"Erreur conversion CSV vers TXT: {e}")
            return False
    
    # Conversions Images
    async def _convert_jpg_to_png(self, source_path: Path, target_path: Path) -> bool:
        try:
            with Image.open(source_path) as img:
                img.save(target_path, 'PNG')
            return True
        except Exception as e:
            print(f"Erreur conversion JPG vers PNG: {e}")
            return False
    
    async def _convert_jpeg_to_png(self, source_path: Path, target_path: Path) -> bool:
        return await self._convert_jpg_to_png(source_path, target_path)
    
    async def _convert_png_to_jpg(self, source_path: Path, target_path: Path) -> bool:
        temp_path = None
        try:
            with Image.open(source_path) as img:
                if img.mode in ('RGBA', 'LA', 'P'):
                    rgb_img = Image.new('RGB', img.size, (255, 255, 255))
                    if img.mode == 'P':
                        img = img.convert('RGBA')
                    rgb_img.paste(img, mask=img.split()[-1] if img.mode == 'RGBA' else None)
                    temp_path = source_path.parent / f"{source_path.stem}_temp.jpg"
                    rgb_img.save(temp_path, 'JPEG', quality=95)
                    os.replace(temp_path, target_path)
                else:
                    img.convert('RGB').save(target_path, 'JPEG', quality=95)
            return True
        except Exception as e:
            print(f"Erreur conversion PNG vers JPG: {e}")
            if temp_path and temp_path.exists():
                temp_path.unlink()
            return False
    
    async def _convert_png_to_jpeg(self, source_path: Path, target_path: Path) -> bool:
        return await self._convert_png_to_jpg(source_path, target_path)

    async def _convert_jpg_to_jpeg(self, source_path: Path, target_path: Path) -> bool:
        try:
            with Image.open(source_path) as img:
                img.convert('RGB').save(target_path, 'JPEG', quality=95)
            return True
        except Exception as e:
            print(f"Erreur conversion JPG vers JPEG: {e}")
            return False

    async def _convert_jpeg_to_jpg(self, source_path: Path, target_path: Path) -> bool:
        return await self._convert_jpg_to_jpeg(source_path, target_path)

    async def _convert_bmp_to_png(self, source_path: Path, target_path: Path) -> bool:
        try:
            with Image.open(source_path) as img:
                # Convert to RGBA if the image has transparency
                if img.mode == 'RGBA':
                    img.save(target_path, 'PNG')
                else:
                    # Convert to RGB for standard images
                    img.convert('RGB').save(target_path, 'PNG')
            return True
        except Exception as e:
            print(f"Erreur conversion BMP vers PNG: {e}")
            return False

    async def _convert_bmp_to_jpg(self, source_path: Path, target_path: Path) -> bool:
        try:
            with Image.open(source_path) as img:
                # Always convert to RGB for JPG
                img = img.convert('RGB')
                img.save(target_path, 'JPEG', quality=95, optimize=True)
            return True
        except Exception as e:
            print(f"Erreur conversion BMP vers JPG: {e}")
            return False
    
    async def _convert_bmp_to_jpeg(self, source_path: Path, target_path: Path) -> bool:
        return await self._convert_bmp_to_jpg(source_path, target_path)

    async def _convert_bmp_to_pdf(self, source_path: Path, target_path: Path) -> bool:
        temp_path = None
        try:
            with Image.open(source_path) as img:
                # Convert to RGB for PDF compatibility
                img = img.convert('RGB')
                temp_path = source_path.parent / f"{source_path.stem}_temp.jpg"
                img.save(temp_path, 'JPEG', quality=95, optimize=True)
                
                with open(target_path, "wb") as f:
                    f.write(img2pdf.convert(str(temp_path)))
                
                if temp_path and temp_path.exists():
                    temp_path.unlink()
            return True
        except Exception as e:
            print(f"Erreur conversion BMP vers PDF: {e}")
            if temp_path and temp_path.exists():
                temp_path.unlink()
            return False
    
    # Conversions TXT
    async def _convert_txt_to_docx(self, source_path: Path, target_path: Path) -> bool:
        try:
            doc = Document()
            
            with open(source_path, 'r', encoding='utf-8') as f:
                content = f.read()
                paragraphs = content.split('\n')
                
                for paragraph in paragraphs:
                    doc.add_paragraph(paragraph)
            
            doc.save(target_path)
            return True
        except Exception as e:
            print(f"Erreur conversion TXT vers DOCX: {e}")
            return False
    
    def get_supported_formats(self) -> dict:
        """Retourne les formats supportés"""
        return self.supported_conversions
